from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

def add_bold_para(text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p

def add_para(text, size=10.5):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(size)
    return p

def add_gray_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p

def add_list_item(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

# ── Header ──
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = h.add_run('屈杰')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('18791410621 | c100_0911@163.com | 21岁 | 应届生 | AI 应用开发工程师').font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GitHub: github.com/echo804')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

# ── 教育背景 ──
add_heading_styled('教育背景', 2)
add_bold_para('西安交通大学城市学院 · 计算机科学与技术 · 本科 · 2023.09 - 2027.06')
add_gray_para('大三在读。大二开始自学 Python 后端和大模型应用开发，线上课程 + 项目驱动学习。能独立阅读英文技术文档（LangChain / PyMuPDF / React 官方文档）。')

# ── 技术能力 ──
add_heading_styled('技术能力', 2)
add_list_item('后端开发：Python / FastAPI / SQLAlchemy async / Pydantic / JWT / RESTful API')
add_list_item('大模型应用：LangChain（LCEL + Prompt Template + ChatOpenAI）/ RAG（pgvector + text-embedding-v3）/ Function Calling / ReAct Agent')
add_list_item('模型对接：OpenAI / DeepSeek / 阿里云百炼——均通过统一接口接入，实测可用')
add_list_item('数据库：PostgreSQL + pgvector / Redis / Alembic')
add_list_item('前端：React 19 + TypeScript + Tailwind CSS v4 + Vite——能独立搭建和开发')
add_list_item('工程化：Docker Compose / Git / Ruff / 模块化分层架构')
add_list_item('文档处理：PyMuPDF / python-docx')

# ── 项目经历 ──
add_heading_styled('项目经历', 2)

add_bold_para('HireMind — AI 智能面试官平台')
add_gray_para('2026.01 - 2026.07 | github.com/echo804/HireMind | 个人项目，独立全栈开发')
add_para('核心功能：上传简历 → AI 自动提取技能/经历 → 生成定制面试题 → 多轮对话面试 → AI 评分报告。')
add_bold_para('我具体做了什么：', 10.5)
add_list_item('多模型面试引擎（~200行核心代码）— 用 LangChain ChatPromptTemplate + ChatOpenAI 构建管道；5 阶段面试大纲（项目深挖→专业考察→问题解决→生产场景→综合设计），LLM 按题号自动匹配；统一接口适配 OpenAI/DeepSeek/百炼。')
add_list_item('RAG 知识检索（~150行核心代码）— PDF/DOCX → 固定窗口切片 → text-embedding-v3 → pgvector；面试中实时语义检索 TOP-3 片段注入 Prompt。')
add_list_item('AI 简历解析 — LLM 提取结构化字段；异步处理 + 百分比进度反馈（解析30% → AI分析50% → 生成报告100%）。')
add_list_item('全栈架构 — models→schemas→repository→service→router 分层；React 19 + Tailwind CSS v4。')
add_bold_para('踩过的坑：', 10.5)
add_list_item('uvicorn --reload 杀掉 asyncio.create_task 后台任务 → 改为上传秒返 + 详情页同步触发分析')
add_list_item('WSL 双副本反复导致"改了代码没生效" → 写了同步脚本')
add_list_item('API Key 脱敏后保存会覆盖真实 Key → 增加掩码检测逻辑')

add_bold_para('QualiGuard — AI 代码质量分析工具')
add_gray_para('2025.07 - 2026.06 | github.com/echo804/QualiGuard | 个人项目')
add_para('核心功能：CLI 工具，自然语言指令驱动 LLM 自动扫描 Python 代码 → 发现问题 → 自动修复 → 生成报告。')
add_bold_para('我具体做了什么：', 10.5)
add_list_item('ReAct Agent 自主执行（~300行）— Thought-Action-Observation 循环；6 个核心能力封装为 OpenAI Function Calling 工具；用户一句话即可完成完整工作流。')
add_list_item('量化评测系统（~200行）— 50 个测试场景覆盖 6 类；指标：通过率、步数、Token 消耗、工具调用准确率；一键切换模型对比。')
add_list_item('多模型适配 — 基于 OpenAI SDK 统一接口，兼容 DeepSeek/Qwen/GLM，仅换 base_url + model 即可切换。')

# ── 自我评价 ──
add_heading_styled('自我评价', 2)
add_list_item('全栈交付能力：独立完成 HireMind 全栈项目（React + FastAPI + PostgreSQL），具备从数据库设计到前端 UI 的完整交付能力。')
add_list_item('大模型应用落地：能搭建从 Prompt 设计到 RAG 检索到多模型适配的完整 AI 应用链路，非纸上谈兵。')
add_list_item('问题解决实例：遇到过 asyncio 后台任务被 reload 杀死、WSL 双副本同步、pgvector 索引参数调优等问题，均独立排查并解决。')
add_list_item('自驱学习：大二起自学 Python 后端 + LangChain + React，通过项目实践驱动学习，能独立查阅英文文档。')

doc.save('/mnt/e/jianli/屈杰-简历.docx')
print('OK')
