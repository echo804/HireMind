from fastapi import APIRouter, Depends, UploadFile, File, Query
from sqlalchemy.ext.asyncio import AsyncSession
from pydantic import BaseModel
import uuid
import logging

logger = logging.getLogger(__name__)

from app.common.result import Result
from app.common.exception.error_code import ErrorCode
from app.common.exception.handlers import BusinessException
from app.common.auth.deps import get_current_user_dev
from app.infrastructure.database import get_db
from app.modules.resume.schemas import ResumeResponse, ResumeListItem, ResumeDetail, UpdateResumeRequest, BatchDeleteRequest
from app.modules.resume.service import ResumeService

router = APIRouter(prefix="/api/resumes", tags=["Resumes"])


@router.post("/upload")
async def upload_resume(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    user_id: uuid.UUID = Depends(get_current_user_dev),
) -> Result[ResumeResponse]:
    service = ResumeService(db)
    content = await file.read()
    result = await service.upload(str(user_id), content, file.filename or "resume")
    return Result.success(result)


@router.get("")
async def list_resumes(
    db: AsyncSession = Depends(get_db),
    user_id: uuid.UUID = Depends(get_current_user_dev),
    q: str | None = Query(None, description="搜索关键词（姓名/职位/摘要）"),
) -> Result[list[ResumeListItem]]:
    service = ResumeService(db)
    result = await service.list_resumes(str(user_id), query=q)
    return Result.success(result)


@router.get("/{resume_id}")
async def get_resume(
    resume_id: str,
    db: AsyncSession = Depends(get_db),
    user_id: uuid.UUID = Depends(get_current_user_dev),
) -> Result[ResumeDetail]:
    service = ResumeService(db)
    result = await service.get_by_id(str(user_id), resume_id)
    return Result.success(result)


@router.delete("/{resume_id}")
async def delete_resume(
    resume_id: str,
    db: AsyncSession = Depends(get_db),
    user_id: uuid.UUID = Depends(get_current_user_dev),
) -> Result[None]:
    service = ResumeService(db)
    await service.delete_resume(str(user_id), resume_id)
    return Result.success(None)


@router.post("/batch-delete")
async def batch_delete_resumes(
    req: BatchDeleteRequest,
    db: AsyncSession = Depends(get_db),
    user_id: uuid.UUID = Depends(get_current_user_dev),
) -> Result[dict]:
    service = ResumeService(db)
    deleted = await service.batch_delete_resumes(str(user_id), req.ids)
    return Result.success({"deleted": deleted})


@router.get("/{resume_id}/duplicate")
async def check_resume_duplicate(
    resume_id: str,
    db: AsyncSession = Depends(get_db),
    user_id: uuid.UUID = Depends(get_current_user_dev),
) -> Result[dict | None]:
    service = ResumeService(db)
    result = await service.get_duplicate(str(user_id), resume_id)
    return Result.success(result)


@router.put("/{resume_id}")
async def update_resume(
    resume_id: str, req: UpdateResumeRequest,
    db: AsyncSession = Depends(get_db),
    user_id: uuid.UUID = Depends(get_current_user_dev),
) -> Result[ResumeDetail]:
    """简历人工校正：更新解析结果字段"""
    service = ResumeService(db)
    result = await service.update_resume(str(user_id), resume_id, req)
    return Result.success(result)


class ExportResumeRequest(BaseModel):
    polished_text: str
    format: str = "docx"


@router.post("/{resume_id}/analyze")
async def analyze_resume_quality_endpoint(
    resume_id: str,
    db: AsyncSession = Depends(get_db),
    user_id: uuid.UUID = Depends(get_current_user_dev),
) -> Result[dict]:
    """AI 按面试官手册对简历分层诊断"""
    service = ResumeService(db)
    result = await service.analyze_quality(str(user_id), resume_id)
    return Result.success(result)


@router.post("/{resume_id}/polish")
async def polish_resume_endpoint(
    resume_id: str,
    db: AsyncSession = Depends(get_db),
    user_id: uuid.UUID = Depends(get_current_user_dev),
) -> Result[dict]:
    """AI 润色简历文本"""
    service = ResumeService(db)
    result = await service.polish(str(user_id), resume_id)
    return Result.success(result)


@router.post("/{resume_id}/export")
async def export_resume(
    resume_id: str, req: ExportResumeRequest,
    db: AsyncSession = Depends(get_db),
    user_id: uuid.UUID = Depends(get_current_user_dev),
):
    """导出 AI 润色后的简历（docx 或 pdf），基于原简历模板排版，生成新文件不覆盖原文件"""
    from fastapi.responses import StreamingResponse
    from io import BytesIO

    service = ResumeService(db)
    detail = await service.get_by_id(str(user_id), resume_id)
    text = (req.polished_text or "").strip()
    if not text:
        return Result.error(ErrorCode.RESUME_NOT_FOUND, "润色内容为空，无法导出")
    original = await service.get_original_path(str(user_id), resume_id)
    filename = (detail.name or "resume").replace(" ", "_") + "_polished"
    from urllib.parse import quote
    encoded = quote(filename)

    if req.format == "pdf":
        return _export_pdf(text, filename, original, detail, encoded)
    return _export_docx(text, filename, original, detail, encoded)


def _export_docx(text: str, filename: str, original, detail, encoded: str):
    """导出 docx：docx 源用原模板重建，pdf 源按 markdown 层级重建"""
    from fastapi.responses import StreamingResponse
    from io import BytesIO
    try:
        from docx import Document

        original_path, orig_type = (original or (None, None))
        doc = None
        if original_path and orig_type == "docx":
            # 以原 docx 为模板：保留页面设置/样式表，清空正文后重建
            doc = Document(original_path)
            for para in list(doc.paragraphs):
                p = para._element
                p.getparent().remove(p)
        else:
            doc = Document()

        for line in text.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=2)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=1)
            elif line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=0)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(line)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}.docx"})
    except ImportError:
        return Result.error(ErrorCode.INTERNAL_ERROR, "docx 支持未安装")
    except Exception as e:
        logger.error(f"export docx failed: {e}")
        return Result.error(ErrorCode.INTERNAL_ERROR, f"导出失败：{str(e)[:100]}")


def _export_pdf(text: str, filename: str, original, detail, encoded: str):
    """导出 pdf：pdf 源复用原版式骨架重绘，docx 源/无源按 markdown 层级生成"""
    from fastapi.responses import StreamingResponse
    from io import BytesIO
    try:
        import fitz  # PyMuPDF

        original_path, orig_type = (original or (None, None))
        doc = fitz.open()
        page = doc.new_page()

        # 基础排版参数
        margin_x = 50
        margin_y = 50
        page_width = page.rect.width
        page_height = page.rect.height
        max_width = page_width - margin_x * 2
        y = margin_y
        font_body = 11
        font_title = 16
        font_sub = 13
        line_height = 16
        title_gap = 4

        # 复用原 pdf 版式骨架：提取原文件首页的字体大小分布
        body_size = font_body
        title_size = font_title
        if original_path and orig_type == "pdf":
            try:
                src = fitz.open(original_path)
                if src.page_count > 0:
                    p0 = src[0]
                    sizes = []
                    for block in p0.get_text("dict")["blocks"]:
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                if span["text"].strip():
                                    sizes.append(span["size"])
                    if sizes:
                        sizes.sort()
                        body_size = sizes[len(sizes) // 2]  # 中位数作正文
                        title_size = max(sizes[-1], body_size + 3)
                src.close()
            except Exception as e:
                logger.warning(f"read pdf skeleton failed: {e}")

        for raw_line in text.split("\n"):
            line = raw_line.strip()
            if not line:
                y += line_height * 0.5
                continue
            if line.startswith("### "):
                size = max(body_size + 2, 12)
                line = line[4:].strip()
            elif line.startswith("## "):
                size = title_size
                line = line[3:].strip()
            elif line.startswith("# "):
                size = title_size + 3
                line = line[2:].strip()
            elif line.startswith("- ") or line.startswith("* "):
                size = body_size
                line = "• " + line[2:].strip()
            else:
                size = body_size

            # 按宽度自动换行
            words = line.split(" ")
            current = ""
            for w in words:
                trial = (current + " " + w).strip()
                tw = fitz.get_text_length(trial, fontname="china-s", fontsize=size)
                if tw > max_width and current:
                    _insert_pdf_line(page, current, margin_x, y, size)
                    y += size + line_height * 0.4
                    current = w
                else:
                    current = trial
            if current:
                _insert_pdf_line(page, current, margin_x, y, size)
                y += size + line_height * 0.4

            if y > page_height - margin_y:
                page = doc.new_page()
                y = margin_y

        buf = BytesIO()
        doc.save(buf)
        doc.close()
        buf.seek(0)
        return StreamingResponse(
            buf, media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}.pdf"})
    except ImportError:
        return Result.error(ErrorCode.INTERNAL_ERROR, "pdf 支持未安装")
    except Exception as e:
        logger.error(f"export pdf failed: {e}")
        return Result.error(ErrorCode.INTERNAL_ERROR, f"导出失败：{str(e)[:100]}")


def _insert_pdf_line(page, text: str, x: float, y: float, size: float):
    """在 pdf 页面插入一行文本（优先中文字体）"""
    try:
        page.insert_text((x, y), text, fontsize=size, fontname="china-s")
    except Exception:
        try:
            page.insert_text((x, y), text, fontsize=size, fontname="helv")
        except Exception:
            pass
