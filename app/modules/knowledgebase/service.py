"""Knowledge base service: parsing, chunking, embedding, search"""

import io, uuid, re, logging, asyncio, os
from typing import Optional
import fitz
from docx import Document as DocxDocument
from sqlalchemy import select, delete, text
from sqlalchemy.ext.asyncio import AsyncSession
from app.common.exception.error_code import ErrorCode
from app.common.exception.handlers import BusinessException
from app.modules.knowledgebase.models import KnowledgeDocument, KnowledgeChunk
from app.infrastructure.cache import cache_get, cache_set, invalidate_user_cache
import hashlib

logger = logging.getLogger(__name__)
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50

# 公共默认文档的归属账号（未登录 DEV 模式上传的文档对所有用户可见）
DEV_USER_ID = uuid.UUID("00000000-0000-0000-0000-000000000000")


class KnowledgeService:
    def __init__(self, db: AsyncSession):
        self.db = db

    def _parse_pdf(self, content: bytes) -> str:
        doc = fitz.open(stream=content, filetype="pdf")
        return "\n\n".join(page.get_text() for page in doc)

    def _parse_docx(self, content: bytes) -> str:
        doc = DocxDocument(io.BytesIO(content))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def _parse_txt(self, content: bytes) -> str:
        return content.decode("utf-8", errors="replace")

    def _parse_document(self, content: bytes, file_type: str) -> str:
        parsers = {"pdf": self._parse_pdf, "docx": self._parse_docx, "txt": self._parse_txt, "md": self._parse_txt}
        parser = parsers.get(file_type)
        if not parser:
            raise BusinessException(ErrorCode.BAD_REQUEST, f"Unsupported format: {file_type}")
        return parser(content)

    def _chunk_text(self, text: str) -> list[str]:
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        sentences = []
        for p in paragraphs:
            parts = re.split(r"(?<=[\u3002\uff01\uff1f.!?\n])", p)
            sentences.extend(s.strip() for s in parts if s.strip())
        chunks, current = [], ""
        for s in sentences:
            if len(current) + len(s) <= CHUNK_SIZE:
                current += s
            else:
                if current:
                    chunks.append(current)
                overlap = current[-CHUNK_OVERLAP:] if current and len(current) > CHUNK_OVERLAP else ""
                current = overlap + s
        if current:
            chunks.append(current)
        return chunks if chunks else [text]

    async def _get_api_config(self, user_id: str | None = None):
        from app.modules.settings.service import get_active_config
        cfg = get_active_config(user_id)
        api_key, base_url = cfg.get("api_key", ""), cfg.get("base_url", "")
        if not api_key:
            raise BusinessException(ErrorCode.BAD_REQUEST, "Configure AI API Key first")
        return api_key, base_url

    async def _embed_chunks(self, chunks: list[str], user_id: str | None = None) -> list[list[float]]:
        import httpx
        api_key, base_url = await self._get_api_config(user_id)
        BATCH = 10
        all_results = []
        async with httpx.AsyncClient(timeout=120) as client:
            for start in range(0, len(chunks), BATCH):
                batch = chunks[start:start + BATCH]
                resp = await client.post(f"{base_url}/embeddings", json={"model": "text-embedding-v3", "input": batch},
                    headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"})
                if resp.status_code != 200:
                    logger.error(f"Embedding API: {resp.status_code} {resp.text[:300]}")
                    raise BusinessException(ErrorCode.AI_SERVICE_UNAVAILABLE, f"Embedding failed: {resp.status_code}")
                data = resp.json()
                result = [None] * len(batch)
                for item in data["data"]:
                    result[item["index"]] = item["embedding"]
                all_results.extend(result)
        return all_results

    async def upload_document(self, user_id: str, content: bytes, filename: str) -> dict:
        ft = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
        if ft not in ("pdf", "docx", "txt", "md"):
            raise BusinessException(ErrorCode.BAD_REQUEST, f"Unsupported: {ft}")
        category = self._guess_category(filename)
        doc = KnowledgeDocument(user_id=uuid.UUID(user_id), filename=filename, file_type=ft, category=category, file_size=len(content), status="processing")
        self.db.add(doc)
        await self.db.commit()
        await self.db.refresh(doc)
        # 落盘原始文件，供后台处理与失败重试复用
        from app.config.settings import settings as s
        raw_path = self._save_raw(s.STORAGE_PATH, str(doc.id), content)
        asyncio.create_task(self._process_document(str(doc.id), raw_path, ft))
        await invalidate_user_cache("kb", user_id)
        return doc.to_dict()

    def _save_raw(self, storage_path: str, doc_id: str, content: bytes) -> str:
        kb_dir = os.path.join(storage_path, "knowledge")
        os.makedirs(kb_dir, exist_ok=True)
        path = os.path.join(kb_dir, f"{doc_id}.bin")
        with open(path, "wb") as f:
            f.write(content)
        return path

    @staticmethod
    def _guess_category(filename: str) -> str:
        # 文件名格式：{主题}_{主题}_{内容}.md（如 llm_llm_decoding_strategies.md），取第二段前缀
        prefix = filename.split("_", 2)[1].lower() if filename.count("_") >= 1 else filename.split("_", 1)[0].lower()
        if prefix not in ("agent", "llm", "rag", "tools", "overview"):
            p1 = filename.split("_", 1)[0].lower()
            prefix = p1 if p1 in ("agent", "llm", "rag", "tools", "overview") else "other"
        return prefix

    async def _process_document(self, doc_id: str, raw_path: str, ft: str):
        """后台处理：解析 → 切片 → 向量化 → 落库（独立 session）"""
        from app.infrastructure.database import async_session_factory
        from sqlalchemy import select as _select
        try:
            async with async_session_factory() as session:
                r = await session.execute(_select(KnowledgeDocument).where(KnowledgeDocument.id == uuid.UUID(doc_id)))
                doc = r.scalar_one_or_none()
                if not doc:
                    return
                doc.status, doc.error_message = "processing", None
                await session.commit()

                with open(raw_path, "rb") as f:
                    content = f.read()
                raw = self._parse_document(content, ft)
                if not raw.strip():
                    raise BusinessException(ErrorCode.BAD_REQUEST, "No extractable text")
                chunks = self._chunk_text(raw)
                logger.info(f"Doc {doc.id}: {len(chunks)} chunks")
                embeddings = await self._embed_chunks(chunks, str(doc.user_id))

                # 删除旧切片（重试场景），写入新切片
                await session.execute(delete(KnowledgeChunk).where(KnowledgeChunk.document_id == doc.id))
                for i, (ct, em) in enumerate(zip(chunks, embeddings)):
                    session.add(KnowledgeChunk(document_id=doc.id, chunk_index=i, content=ct, embedding=em))
                doc.status, doc.chunk_count = "ready", len(chunks)
                await session.commit()
                await invalidate_user_cache("kb", str(doc.user_id))
        except Exception as e:
            logger.error(f"Process doc {doc_id} failed: {e}")
            try:
                async with async_session_factory() as session:
                    r = await session.execute(_select(KnowledgeDocument).where(KnowledgeDocument.id == uuid.UUID(doc_id)))
                    doc = r.scalar_one_or_none()
                    if doc:
                        doc.status, doc.error_message = "failed", str(e)
                        await session.commit()
            except Exception as e2:
                logger.error(f"Mark doc {doc_id} failed: {e2}")

    async def get_document_status(self, doc_id: str, user_id: str) -> dict:
        did, uid = uuid.UUID(doc_id), uuid.UUID(user_id)
        r = await self.db.execute(
            select(KnowledgeDocument).where(
                KnowledgeDocument.id == did,
                KnowledgeDocument.user_id.in_([uid, DEV_USER_ID]),
            )
        )
        doc = r.scalar_one_or_none()
        if not doc:
            raise BusinessException(ErrorCode.NOT_FOUND, "Document not found")
        return doc.to_dict()

    async def retry_document(self, doc_id: str, user_id: str) -> dict:
        """失败重试：读取落盘原始文件重新处理"""
        from app.config.settings import settings as s
        did, uid = uuid.UUID(doc_id), uuid.UUID(user_id)
        r = await self.db.execute(
            select(KnowledgeDocument).where(
                KnowledgeDocument.id == did,
                KnowledgeDocument.user_id.in_([uid, DEV_USER_ID]),
            )
        )
        doc = r.scalar_one_or_none()
        if not doc:
            raise BusinessException(ErrorCode.NOT_FOUND, "Document not found")
        if doc.status == "processing":
            return doc.to_dict()
        raw_path = os.path.join(s.STORAGE_PATH, "knowledge", f"{doc_id}.bin")
        if not os.path.exists(raw_path):
            raise BusinessException(ErrorCode.BAD_REQUEST, "原始文件不存在，请重新上传")
        doc.status, doc.error_message = "processing", None
        await self.db.commit()
        asyncio.create_task(self._process_document(str(doc.id), raw_path, doc.file_type))
        await invalidate_user_cache("kb", user_id)
        return doc.to_dict()




    async def get_document_content(self, doc_id: str, user_id: str) -> dict:
        """Get document with all its chunks"""
        did, uid = uuid.UUID(doc_id), uuid.UUID(user_id)
        r = await self.db.execute(
            select(KnowledgeDocument).where(
                KnowledgeDocument.id == did,
                KnowledgeDocument.user_id.in_([uid, DEV_USER_ID]),
            )
        )
        doc = r.scalar_one_or_none()
        if not doc:
            raise BusinessException(ErrorCode.NOT_FOUND, "Document not found")

        # Get chunks
        cr = await self.db.execute(
            select(KnowledgeChunk)
            .where(KnowledgeChunk.document_id == did)
            .order_by(KnowledgeChunk.chunk_index)
        )
        chunks = cr.scalars().all()

        return {
            **doc.to_dict(),
            "chunks": [
                {"index": c.chunk_index, "content": c.content}
                for c in chunks
            ],
        }

    async def list_documents(self, user_id: str) -> list[dict]:
        cached = await cache_get("kb", "list", user_id)
        if cached is not None:
            return cached
        r = await self.db.execute(select(KnowledgeDocument).where(KnowledgeDocument.user_id.in_([uuid.UUID(user_id), DEV_USER_ID])).order_by(KnowledgeDocument.created_at.desc()))
        result = [d.to_dict() for d in r.scalars().all()]
        await cache_set("kb", "list", user_id, data=result, ttl=600)
        return result

    async def delete_document(self, doc_id: str, user_id: str):
        did, uid = uuid.UUID(doc_id), uuid.UUID(user_id)
        r = await self.db.execute(select(KnowledgeDocument).where(KnowledgeDocument.id == did, KnowledgeDocument.user_id == uid))
        doc = r.scalar_one_or_none()
        if not doc:
            raise BusinessException(ErrorCode.NOT_FOUND, "Document not found")
        await self.db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.document_id == did))
        await self.db.delete(doc)
        await self.db.commit()
        await invalidate_user_cache("kb", user_id)

    async def search(self, query: str, top_k: int = 3, user_id: Optional[str] = None) -> list[dict]:
        import httpx
        api_key, base_url = await self._get_api_config(user_id)
        async with httpx.AsyncClient(timeout=30) as client:
            resp = await client.post(f"{base_url}/embeddings", json={"model": "text-embedding-v3", "input": query},
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"})
            if resp.status_code != 200:
                logger.error(f"Query embedding: {resp.status_code} {resp.text[:300]}")
                raise BusinessException(ErrorCode.AI_SERVICE_UNAVAILABLE, f"Query embedding: {resp.status_code}")
            qvec = resp.json()["data"][0]["embedding"]
        vs = "[" + ",".join(str(v) for v in qvec) + "]"
        # 缓存 key 按 user_id 隔离，先读缓存
        query_hash = hashlib.md5(f"{query}:{top_k}".encode()).hexdigest()[:12]
        cache_key = f"{user_id or 'anon'}:{query_hash}"
        cached = await cache_get("kb", "search", cache_key)
        if cached is not None:
            return cached
        # 搜索范围：自己的文档 + 公共默认文档（DEV 账号）
        wc = "AND (kd.user_id = CAST(:uid AS uuid) OR kd.user_id = CAST(:dev AS uuid))" if user_id else ""
        params = {"emb": vs, "emb2": vs, "top": top_k}
        if user_id:
            params["uid"] = user_id
            params["dev"] = str(DEV_USER_ID)
        sql = text(f"""
            SELECT kc.content, kc.chunk_index,
                   1 - (kc.embedding <=> CAST(:emb AS vector)) AS score,
                   kd.filename AS doc_name,
                   kc.id AS chunk_id,
                   kd.id AS document_id
            FROM knowledge_chunks kc
            JOIN knowledge_documents kd ON kd.id = kc.document_id
            WHERE kd.status = 'ready' {wc}
            ORDER BY kc.embedding <=> CAST(:emb2 AS vector)
            LIMIT :top
        """)
        rows = (await self.db.execute(sql, params)).fetchall()
        result = [{"content": r[0], "chunk_index": r[1], "score": float(r[2]), "document_name": r[3], "chunk_id": str(r[4]), "document_id": str(r[5])} for r in rows]
        await cache_set("kb", "search", cache_key, data=result, ttl=300)
        return result

    async def qa_answer(self, user_id: str, question: str, top_k: int = 3) -> dict:
        """基于知识库检索的问答：检索相关切片 → 拼上下文 → LLM 生成答案（标注来源）"""
        results = await self.search(question, top_k, user_id)
        if not results:
            return {
                "answer": "知识库中暂时没有与这个问题相关的内容。请先上传相关文档，或换个问法试试。",
                "sources": [],
            }

        from app.config.settings import settings as s
        from app.modules.interview.agent import _get_llm
        llm = _get_llm(s, user_id)
        if not llm:
            return {
                "answer": "AI 服务未配置或不可用，请先在系统设置中配置有效的 API Key。",
                "sources": [
                    {"document_name": r["document_name"], "chunk_index": r["chunk_index"], "content": r["content"], "score": r["score"], "document_id": r["document_id"], "chunk_id": r["chunk_id"]}
                    for r in results
                ],
            }

        # 拼接检索到的切片作为上下文
        context_parts = []
        for i, r in enumerate(results, 1):
            context_parts.append(f"[片段{i} 来自《{r['document_name']}》]\n{r['content']}")
        context = "\n\n".join(context_parts)
        doc_names = "、".join(dict.fromkeys(r["document_name"] for r in results))

        prompt = f"""你是一名知识库问答助手。请仅根据下面提供的知识库内容回答用户问题。

## 知识库内容
{context}

## 用户问题
{question}

回答要求：
1. 只依据上述知识库内容回答，不要编造知识库中不存在的信息；若内容不足以回答，如实说明
2. 回答用中文，条理清晰，可适当分点
3. 回答末尾标注引用的来源：\u3010来源: {doc_names}\u3011
"""
        try:
            resp = await llm.ainvoke(prompt)
            answer = resp.content if hasattr(resp, "content") else str(resp)
        except Exception as e:
            logger.error(f"QA LLM call failed: {e}")
            answer = f"AI 生成回答时出错：{e}"

        return {
            "answer": answer,
            "sources": [
                {"document_name": r["document_name"], "chunk_index": r["chunk_index"], "content": r["content"], "score": r["score"], "document_id": r["document_id"], "chunk_id": r["chunk_id"]}
                for r in results
            ],
        }
